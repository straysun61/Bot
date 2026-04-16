"""
核心 RAG Engine 模块
负责文档解析、向量化和检索
"""
import os
import re
import json
import base64
import uuid
import tempfile
from typing import Optional, Tuple, List, Dict
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain.retrievers.multi_vector import MultiVectorRetriever
from langchain_core.stores import InMemoryStore
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings
import httpx

from core.config import settings
import logging
import pickle

logger = logging.getLogger(__name__)


class RerankerClient:
    """
    Reranker 客户端，调用远程 rerank 服务进行二阶段重排
    """

    def __init__(self, reranker_url: str):
        self.reranker_url = reranker_url.rstrip("/")
        self.rerank_endpoint = f"{self.reranker_url}/rerank"

    def rerank(self, query: str, documents: list[Document], top_n: int = 5) -> list[Document]:
        """
        对文档列表进行重排

        Args:
            query: 查询文本
            documents: RRF 融合后的候选文档列表
            top_n: 返回的 top n 结果

        Returns:
            重排后的文档列表
        """
        if not documents:
            return []

        doc_texts = [doc.page_content for doc in documents]

        try:
            import httpx
            response = httpx.post(
                self.rerank_endpoint,
                json={
                    "query": query,
                    "documents": doc_texts,
                    "top_n": top_n
                },
                timeout=60.0
            )
            response.raise_for_status()
            result = response.json()

            results = result.get("results", [])
            # 按 relevance_score 降序排序
            results.sort(key=lambda x: x.get("relevance_score", 0), reverse=True)

            reranked = []
            for item in results[:top_n]:
                idx = item.get("index")
                if idx is not None and 0 <= idx < len(documents):
                    reranked.append(documents[idx])

            logger.info(f"[RERANK] 重排完成，输入 {len(documents)} 条，输出 {len(reranked)} 条")
            return reranked

        except Exception as e:
            logger.warning(f"[RERANK] 重排失败: {e}，返回 RRF 结果")
            return documents[:top_n]


class BM25Retriever:
    """
    BM25 关键词检索器，支持持久化到磁盘
    """

    def __init__(self, persist_path: str):
        self.persist_path = persist_path
        self.doc_store: list[tuple[str, str, dict]] = []  # [(doc_id, text, metadata), ...]
        self.index = None
        self._tokenized_corpus: list[list[str]] = []
        self._loaded = False

    def _tokenize(self, text: str) -> list[str]:
        """简单分词（按字符级别，适合中文）"""
        # 简单按标点和空格分词
        import re
        tokens = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z0-9]+', text)
        return [t.lower() for t in tokens if len(t) > 1]

    def add_documents(self, documents: list[Document]) -> None:
        """添加文档到 BM25 索引"""
        from rank_bm25 import BM25L

        texts = [doc.page_content for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        doc_ids = [meta.get("doc_id", "") or meta.get("parent_id", "") for meta in metadatas]

        # 追加到 doc_store
        for doc_id, text, meta in zip(doc_ids, texts, metadatas):
            self.doc_store.append((doc_id, text, meta))

        # 重新构建索引
        tokenized = [self._tokenize(text) for text in texts]
        self._tokenized_corpus = tokenized

        if tokenized:
            self.index = BM25L(tokenized)
        else:
            self.index = None

    def save(self) -> None:
        """持久化索引到磁盘"""
        import os
        os.makedirs(os.path.dirname(self.persist_path) or ".", exist_ok=True)
        with open(self.persist_path, "wb") as f:
            pickle.dump({
                "doc_store": self.doc_store,
                "tokenized_corpus": self._tokenized_corpus,
            }, f)
        logger.info(f"[BM25] 索引已保存到 {self.persist_path}，共 {len(self.doc_store)} 条")

    def load(self) -> None:
        """从磁盘加载索引"""
        import os
        if not os.path.exists(self.persist_path):
            logger.info("[BM25] 索引文件不存在，将从头构建")
            return

        try:
            with open(self.persist_path, "rb") as f:
                data = pickle.load(f)
            self.doc_store = data["doc_store"]
            self._tokenized_corpus = data["tokenized_corpus"]

            if self._tokenized_corpus:
                from rank_bm25 import BM25L
                self.index = BM25L(self._tokenized_corpus)
            logger.info(f"[BM25] 索引已加载，共 {len(self.doc_store)} 条")
        except Exception as e:
            logger.warning(f"[BM25] 索引加载失败: {e}，将重新构建")
            self.doc_store = []
            self._tokenized_corpus = []
            self.index = None

    def search(self, query: str, k: int = 5) -> list[Document]:
        """搜索返回 top-k 相关文档"""
        if not self.index or not self._tokenized_corpus:
            return []

        query_tokens = self._tokenize(query)
        scores = self.index.get_scores(query_tokens)

        # 按分数排序，取 top-k
        doc_scores = list(enumerate(scores))
        doc_scores.sort(key=lambda x: x[1], reverse=True)

        results = []
        for idx, score in doc_scores[:k]:
            if score > 0:
                doc_id, text, meta = self.doc_store[idx]
                doc = Document(page_content=text, metadata=meta)
                results.append(doc)

        return results

# 图片处理模块
from core.image_handler import (
    extract_images_from_pdf,
    extract_images_from_docx,
    save_image_records,
    get_image_record_markdown,
    ImageHandler
)


class FallbackToOCRError(Exception):
    """轨道A解析失败，需要回退到轨道B OCR"""
    pass


class DashScopeEmbeddings(Embeddings):
    """DashScope (阿里云百炼) Embedding 实现"""

    def __init__(self, model: str = "text-embedding-v3", api_key: str = None, base_url: str = None):
        self.model = model
        self.api_key = api_key or settings.OPENAI_API_KEY
        self.base_url = base_url or settings.OPENAI_API_BASE

    def _call_api(self, texts: List[str]) -> List[List[float]]:
        """调用 DashScope embedding API，支持单条失败时逐条处理"""
        os.environ["HF_HUB_OFFLINE"] = "1"
        os.environ["TRANSFORMERS_OFFLINE"] = "1"

        try:
            response = httpx.post(
                f"{self.base_url}/embeddings",
                headers={
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                },
                json={"model": self.model, "input": texts},
                timeout=60.0
            )
            response.raise_for_status()
            result = response.json()
            return [item["embedding"] for item in result["data"]]
        except httpx.HTTPStatusError as e:
            # 批量失败时，回退到逐条调用
            if e.response.status_code == 400 and len(texts) > 1:
                logger.info(f"[Embedding] 批量请求失败(400)，回退到逐条处理 {len(texts)} 条文本")
                embeddings = []
                for i, text in enumerate(texts):
                    try:
                        single_resp = httpx.post(
                            f"{self.base_url}/embeddings",
                            headers={
                                "Authorization": f"Bearer {self.api_key}",
                                "Content-Type": "application/json"
                            },
                            json={"model": self.model, "input": [text]},
                            timeout=30.0
                        )
                        single_resp.raise_for_status()
                        emb = single_resp.json()["data"][0]["embedding"]
                        embeddings.append(emb)
                        logger.warning(f"[EMBED] Text Embedding] 第 {i+1}/{len(texts)} 条成功")
                    except Exception as ex:
                        logger.warning(f"[EMBED] Text Embedding] 第 {i+1}/{len(texts)} 条失败: {ex}，使用模拟向量")
                        embeddings.append(self._mock_embedding(text))
                return embeddings
            raise

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """批量嵌入文档，支持分批和重试"""
        # 分批发送，避免单次请求过大
        batch_size = 10
        all_embeddings = []
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            try:
                batch_embeddings = self._call_api(batch)
                all_embeddings.extend(batch_embeddings)
            except Exception as e:
                logger.warning(f"[EMBED] Batch {i//batch_size} failed: {e}")
                # 尝试逐条发送，定位问题
                for j, text in enumerate(batch):
                    try:
                        emb = self._call_api([text])
                        all_embeddings.extend(emb)
                    except Exception as e2:
                        logger.warning(f"[EMBED] Text EMBED] Text {i+j} failed: {e2}")
                        logger.warning(f"[EMBED] Text EMBED]   Text: {repr(text[:200])}")
                        # 使用零向量作为fallback
                        all_embeddings.append([0.0] * 1024)
        return all_embeddings

    def embed_query(self, text: str) -> List[float]:
        """嵌入单个查询"""
        return self._call_api([text])[0]

    def _mock_embedding(self, text: str) -> List[float]:
        """生成模拟向量（当 API 调用失败时使用）"""
        import hashlib
        hash_bytes = hashlib.md5(text.encode()).digest()
        return [float(b) / 255.0 for b in hash_bytes[:16]] * 8


class ImagePreprocessor:
    """
    优化方案 1: 针对手写理科作业的图像预处理通道

    使用 OpenCV 对图像进行多步提纯：
    - 灰度化 + CLAHE（对比度受限自适应直方图均衡）提亮淡墨水
    - 自适应二值化（高斯/均值）漂白背景
    - 形态学膨胀加深笔画边缘
    - 可选去噪（中值滤波）
    """

    @staticmethod
    def preprocess(image_path: str, strategy: str = "science") -> str:
        """
        图像预处理入口。

        Args:
            image_path: 原始图片路径
            strategy: 预处理策略
                - "science": 理科公式优化（强对比度 + 形态学加粗）
                - "general": 通用文字（温和增强）
                - "faint_ink": 淡墨水增强（极端对比度）

        Returns:
            预处理后的图片路径（临时文件，调用方负责清理）
        """
        try:
            import cv2
        except ImportError:
            logger.warning("OpenCV 未安装，跳过图像预处理")
            return image_path

        img = cv2.imread(image_path)
        if img is None:
            logger.warning(f"无法读取图片: {image_path}")
            return image_path

        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)

        if strategy == "science":
            processed = ImagePreprocessor._science_preprocess(gray)
        elif strategy == "faint_ink":
            processed = ImagePreprocessor._faint_ink_preprocess(gray)
        else:
            processed = ImagePreprocessor._general_preprocess(gray)

        # 保存为临时 PNG 文件
        import tempfile
        temp_file = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
        temp_path = temp_file.name
        temp_file.close()
        cv2.imwrite(temp_path, processed, [cv2.IMWRITE_PNG_COMPRESSION, 3])
        logger.info(f"图像预处理完成 [{strategy}]: {image_path} -> {temp_path}")
        return temp_path

    @staticmethod
    def _clahe_enhance(gray: "np.ndarray", clip_limit: float = 2.0, tile_size: int = 8) -> "np.ndarray":
        """CLAHE 对比度增强，解决局部曝光不均"""
        import numpy as np
        clahe = cv2.createCLAHE(clipLimit=clip_limit, tileGridSize=(tile_size, tile_size))
        return clahe.apply(gray)

    @staticmethod
    def _adaptive_threshold(gray: "np.ndarray", block_size: int = 25, C: int = 10) -> "np.ndarray":
        """自适应二值化 - 高斯加权"""
        return cv2.adaptiveThreshold(
            gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY, block_size, C
        )

    @staticmethod
    def _morphological_bold(binary: "np.ndarray", kernel_size: int = 2) -> "np.ndarray":
        """形态学膨胀：加深笔画边缘，使化学脚标/物理角标更清晰"""
        kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (kernel_size, kernel_size))
        # 先膨胀后腐蚀（闭运算），填补笔画内部空洞
        return cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)

    @staticmethod
    def _denoise(gray: "np.ndarray", kernel: int = 3) -> "np.ndarray":
        """中值滤波去噪，去除椒盐噪声"""
        return cv2.medianBlur(gray, kernel)

    @staticmethod
    def _science_preprocess(gray: "np.ndarray") -> "np.ndarray":
        """理科公式优化流程：CLAHE → 自适应二值化 → 形态学加粗"""
        enhanced = ImagePreprocessor._clahe_enhance(gray, clip_limit=3.0, tile_size=8)
        denoised = ImagePreprocessor._denoise(enhanced, kernel=3)
        binary = ImagePreprocessor._adaptive_threshold(denoised, block_size=31, C=8)
        bold = ImagePreprocessor._morphological_bold(binary, kernel_size=2)
        return bold

    @staticmethod
    def _faint_ink_preprocess(gray: "np.ndarray") -> "np.ndarray":
        """淡墨水极端增强：强力 CLAHE → 自适应二值化 → 粗膨胀"""
        enhanced = ImagePreprocessor._clahe_enhance(gray, clip_limit=5.0, tile_size=16)
        denoised = ImagePreprocessor._denoise(enhanced, kernel=5)
        binary = ImagePreprocessor._adaptive_threshold(denoised, block_size=51, C=5)
        bold = ImagePreprocessor._morphological_bold(binary, kernel_size=3)
        return bold

    @staticmethod
    def _general_preprocess(gray: "np.ndarray") -> "np.ndarray":
        """通用温和增强：轻度 CLAHE → 轻微二值化"""
        enhanced = ImagePreprocessor._clahe_enhance(gray, clip_limit=1.5, tile_size=8)
        return enhanced


class VisionOCR:
    """
    视觉大模型 OCR 处理器 - 使用阿里云百炼 qwen-vl-max
    优化后支持：
    - 图像预处理（CLAHE + 自适应二值化 + 形态学）
    - Few-Shot 多模态注入（化学/物理/数学高质量范例）
    - 双轨降级（VLM → SimpleTex 公式识别）
    """

    # ==================== 核心 System Prompt ====================

    SYSTEM_PROMPT = """你是一个专精于高中理科手写作业 OCR 与公式排版的顶级引擎。
请严格遵守以下规则提取图片中的所有信息：

【排版规则】
1. 绝对忠实原文：如实识别每一行文本，不得总结、删减或编造。
2. 多栏排版：如果有多栏或左右排版，根据语义逻辑转换为从上到下的阅读顺序。
3. 层级结构：遇到"一、计算题""二、填空题"等大纲，使用 Markdown 标题（##/###）标记。

【公式规范 - 严格执行】
4. 所有数学/物理/化学公式统一用 $$ 包裹的标准 LaTeX 格式输出：
   - 分数：$$ \\frac{分子}{分母} $$  （禁止使用 1/20 等纯文本）
   - 幂/下标：$$ v_{0}^{2} $$、$$ F_{N} $$、$$ H_{2}O $$
   - 物理向量：$$ \\vec{F} $$、$$ \\overrightarrow{AB} $$
   - 求和/积分：$$ \\sum_{i=1}^{n} $$、$$ \\int_{a}^{b} $$

【学科特殊规则】
5. 化学方程式：必须用 LaTeX 化学环境或标准公式格式
   - 反应条件写在箭头上方：$$ \\xrightarrow[\\text{加热}]{\\text{MnO}_2} $$
   - 气体/沉淀：$$ \\uparrow $$、$$ \\downarrow $$
   - 化学键：单键 -，双键 =，三键 ≡
   - 示例：$$ 2H_2O_2 \\xrightarrow{MnO_2} 2H_2O + O_2 \\uparrow $$

6. 物理受力分析：
   - 力的表示：$$ F_N $$（支持力）、$$ f $$（摩擦力）、$$ mg $$（重力）
   - 牛顿第二定律：$$ \\sum \\vec{F} = m\\vec{a} $$
   - 动能定理：$$ W = \\Delta E_k = \\frac{1}{2}mv^2 - \\frac{1}{2}mv_0^2 $$

7. 生物：
   - 基因型：$$ AaBb $$、$$ X^B X^b $$
   - 细胞分裂：用箭头表示过程

【手写特有标记处理】
8. 补充插入：箭头指向旁边的补充文字，根据语义插入到正文对应位置。
9. 涂抹/涂改：被完全涂黑的文字忽略，部分涂改用 ~~删除线~~ 标记。
10. 随意勾画：与解题无关的随意涂鸦、划线全部忽略，不输出。"""

    RELAY_SYSTEM_PROMPT = """{base_prompt}

【接力上下文】以下是前一页/图片的内容摘要：
{previous_summary}

请判断当前图片是否为前一页的逻辑延续（如同一个公式的推演、同一段笔记的继续）。
如果是，请保持上下文连贯，继续输出。
如果不是，请独立处理当前内容。

请在输出开头标注 [CONTINUATION] 或 [NEW_START] 以示区分。"""

    MAX_IMAGE_SIZE = 10 * 1024 * 1024

    # ==================== 双轨配置 ====================
    # SimpleTex API 配置（专业公式识别引擎）
    SIMPLETEX_API_URL = "https://simpletex.cn/api/v1/ocr"
    SIMPLETEX_API_KEY = ""  # 在 .env 中配置 SIMPLETEX_API_KEY
    VLM_TIMEOUT_THRESHOLD = 90  # VLM 超时阈值（秒），超过则降级
    FORMULA_DENSITY_THRESHOLD = 0.3  # 公式密度阈值，VLM 输出中公式占比低于此值时触发降级

    @classmethod
    def extract_text_from_image(cls, image_path: str, preprocess: bool = True, strategy: str = "science") -> str:
        """
        从图片中提取文本（入口方法）

        Args:
            image_path: 图片路径
            preprocess: 是否进行图像预处理（默认开启）
            strategy: 预处理策略 ("science" / "faint_ink" / "general")
        """
        if preprocess:
            try:
                processed_path = ImagePreprocessor.preprocess(image_path, strategy)
                if processed_path != image_path:
                    try:
                        result = cls._call_vision(processed_path, cls.SYSTEM_PROMPT)
                    finally:
                        if os.path.exists(processed_path):
                            os.remove(processed_path)
                    return result
            except Exception as e:
                logger.warning(f"图像预处理失败，使用原图: {e}")

        return cls._call_vision(image_path, cls.SYSTEM_PROMPT)

    @classmethod
    async def extract_with_context(
        cls, image_path: str, context: str = "", is_continuation: bool = False,
        preprocess: bool = True, strategy: str = "science"
    ) -> str:
        """带接力上下文的 OCR"""
        prompt = cls.RELAY_SYSTEM_PROMPT.format(
            base_prompt=cls.SYSTEM_PROMPT,
            previous_summary=context
        ) if (is_continuation and context) else cls.SYSTEM_PROMPT

        if preprocess:
            try:
                processed_path = ImagePreprocessor.preprocess(image_path, strategy)
                if processed_path != image_path:
                    try:
                        return cls._call_vision(processed_path, prompt)
                    finally:
                        if os.path.exists(processed_path):
                            os.remove(processed_path)
            except Exception as e:
                logger.warning(f"图像预处理失败，使用原图: {e}")

        return cls._call_vision(image_path, prompt)

    @classmethod
    def _build_messages(cls, image_base64: str, system_prompt: str) -> list:
        """
        优化方案 2: 构建包含 Few-Shot 范例的多模态消息
        """
        # Few-Shot 示例 1: 化学方程式 + 物理公式混合
        example_chemistry_physics = """【输入图片内容】这是一张高中物理+化学作业的手写扫描件，包含：
- 一道物理计算题（关于牛顿第二定律的受力分析）
- 一个化学方程式（实验室制取氧气）

【期望输出】
## 物理计算题

已知物体质量 $$ m = 2 \\text{kg} $$，在水平面上受到拉力 $$ F = 10 \\text{N} $$，摩擦系数 $$ \\mu = 0.3 $$。

### 受力分析
物体受力情况：
- 重力：$$ G = mg = 2 \\times 9.8 = 19.6 \\text{N} $$
- 支持力：$$ F_N = G = 19.6 \\text{N} $$
- 摩擦力：$$ f = \\mu F_N = 0.3 \\times 19.6 = 5.88 \\text{N} $$
- 合外力：$$ F_{合} = F - f = 10 - 5.88 = 4.12 \\text{N} $$

由牛顿第二定律：
$$ a = \\frac{F_{合}}{m} = \\frac{4.12}{2} = 2.06 \\text{m/s}^2 $$

## 化学实验

实验室用双氧水制取氧气：
$$ 2H_2O_2 \\xrightarrow{MnO_2} 2H_2O + O_2 \\uparrow $$

反应条件：二氧化锰催化，常温下即可进行。
收集方法：排水法或向上排空气法。"""

        # Few-Shot 示例 2: 数学推导 + 涂抹处理
        example_math = """【输入图片内容】这是一张数学作业手写件，包含：
- 分数运算推导
- 部分被涂改的字迹
- 旁边有箭头补充

【期望输出】
## 计算题

$$ \\frac{3}{4} + \\frac{1}{6} = \\frac{9}{12} + \\frac{2}{12} = \\frac{11}{12} $$

解方程：
$$ \\frac{x}{3} + \\frac{x}{4} = 7 $$

$$ \\frac{4x + 3x}{12} = 7 $$

$$ \\frac{7x}{12} = 7 $$

$$ x = 12 $$

验证：$$ \\frac{12}{3} + \\frac{12}{4} = 4 + 3 = 7 \\checkmark $$

## 应用题

（被涂改的部分已忽略）

补充：乙单独做需要 $$ 10 $$ 小时完成。
设总工作量为 $$ 1 $$，则甲的工作效率为 $$ \\frac{1}{20} $$，乙的工作效率为 $$ \\frac{1}{10} $$。

两人合作完成时间：
$$ 1 \\div \\left( \\frac{1}{20} + \\frac{1}{10} \\right) = 1 \\div \\frac{3}{20} = \\frac{20}{3} \\text{小时} $$"""

        return [
            {"role": "system", "content": system_prompt},
            {
                "role": "user",
                "content": "以下是一个化学+物理作业的手写图片识别示例，请学习这种输出格式："
            },
            {"role": "assistant", "content": example_chemistry_physics},
            {
                "role": "user",
                "content": "以下是另一个数学作业的手写识别示例，注意公式格式和涂抹处理："
            },
            {"role": "assistant", "content": example_math},
            {
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
                ]
            },
        ]

    @classmethod
    def _call_vision(cls, image_path: str, system_prompt: str, use_few_shot: bool = True) -> str:
        """
        调用视觉大模型进行 OCR

        优化方案 2: 使用 Few-Shot 示例构建消息
        优化方案 3: 内置超时降级到 SimpleTex API
        """
        if not settings.OPENAI_API_KEY:
            return f"[模拟 OCR 结果] 已从图片 {os.path.basename(image_path)} 中提取文本。"

        actual_path = image_path
        file_size = os.path.getsize(image_path)

        if file_size > cls.MAX_IMAGE_SIZE:
            logger.info(f"图片太大 ({file_size / 1024 / 1024:.1f}MB)，正在压缩...")
            actual_path = cls._compress_image(image_path)

        with open(actual_path, "rb") as image_file:
            image_base64 = base64.b64encode(image_file.read()).decode("utf-8")

        # 构建消息（带 Few-Shot）
        messages = cls._build_messages(image_base64, system_prompt) if use_few_shot else [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": [{"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}]},
        ]

        # 优化方案 3: 双轨调用 - 先尝试 VLM，超时则降级到 SimpleTex
        try:
            import concurrent.futures

            def call_vlm():
                response = httpx.post(
                    f"{settings.OPENAI_API_BASE}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {settings.OPENAI_API_KEY}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": "qwen-vl-max",
                        "messages": messages,
                        "max_tokens": 4096
                    },
                    timeout=cls.VLM_TIMEOUT_THRESHOLD
                )
                response.raise_for_status()
                result = response.json()
                return result["choices"][0]["message"]["content"]

            # 在线程池中执行 VLM 调用（带超时控制）
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(call_vlm)
                try:
                    content = future.result(timeout=cls.VLM_TIMEOUT_THRESHOLD + 10)
                    logger.info("VLM OCR 调用成功")
                    return content
                except concurrent.futures.TimeoutError:
                    logger.warning(f"VLM OCR 超时（>{cls.VLM_TIMEOUT_THRESHOLD}s），降级到 SimpleTex...")
                    return cls._fallback_to_sciencetex(actual_path)

        except httpx.HTTPStatusError as e:
            logger.warning(f"VLM OCR HTTP 错误: {e.response.status_code}，降级到 SimpleTex...")
            return cls._fallback_to_sciencetex(actual_path)
        except Exception as e:
            logger.warning(f"VLM OCR 调用异常: {e}，降级到 SimpleTex...")
            return cls._fallback_to_sciencetex(actual_path)
        finally:
            if actual_path != image_path and os.path.exists(actual_path):
                try:
                    os.remove(actual_path)
                except:
                    pass

    @classmethod
    def _fallback_to_sciencetex(cls, image_path: str) -> str:
        """
        优化方案 3: 专业公式识别 API 降级（SimpleTex）
        当 VLM 超时或失败时，使用 SimpleTex 识别公式
        """
        api_key = getattr(settings, 'SIMPLETEX_API_KEY', '') or cls.SIMPLETEX_API_KEY

        if not api_key:
            logger.warning("SimpleTex API 未配置，返回降级提示")
            return "[公式识别服务暂时不可用，请检查配置]"

        try:
            import base64
            with open(image_path, "rb") as f:
                img_b64 = base64.b64encode(f.read()).decode("utf-8")

            response = httpx.post(
                cls.SIMPLETEX_API_URL,
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "image": img_b64,
                    "formula": True,  # 启用公式识别
                    "lang": "zh"
                },
                timeout=30.0
            )
            response.raise_for_status()
            result = response.json()

            if result.get("code") == 200 or result.get("status") == "success":
                markdown = result.get("markdown", result.get("result", ""))
                logger.info(f"SimpleTex 公式识别成功，提取 {len(markdown)} 字符")
                return markdown
            else:
                logger.warning(f"SimpleTex 返回异常: {result}")
                return "[公式识别失败]"

        except httpx.HTTPStatusError as e:
            logger.warning(f"SimpleTex HTTP 错误: {e.response.status_code}")
            return "[公式识别服务暂时不可用]"
        except Exception as e:
            logger.warning(f"SimpleTex 调用异常: {e}")
            return "[公式识别服务暂时不可用]"

    @staticmethod
    def slice_long_image(image_path: str, overlap_px: int = 200, max_height: int = 4096) -> List[str]:
        """
        带重叠的长图切分
        检测图片高度 > max_height 时，按 max_height 步长切分
        相邻切片之间保留 overlap_px 重叠区域

        Returns:
            切片路径列表。如果不需要切分，返回 [image_path]
        """
        from PIL import Image

        img = Image.open(image_path)
        width, height = img.size

        if height <= max_height:
            return [image_path]

        slices = []
        base_name = os.path.splitext(image_path)[0]
        slice_idx = 0
        y_start = 0

        while y_start < height:
            y_end = min(y_start + max_height, height)

            if slice_idx > 0:
                y_start = max(0, y_end - max_height + overlap_px)
                y_end = min(y_start + max_height, height)

            slice_img = img.crop((0, y_start, width, y_end))
            slice_path = f"{base_name}_slice_{slice_idx:03d}.png"
            slice_img.save(slice_path)
            slices.append(slice_path)

            y_start = y_end
            slice_idx += 1

        return slices

    @classmethod
    def _compress_image(cls, image_path: str) -> str:
        """压缩图片使其小于 MAX_IMAGE_SIZE"""
        from PIL import Image
        import tempfile

        img = Image.open(image_path)
        quality = 85
        temp_file = tempfile.NamedTemporaryFile(suffix=".jpg", delete=False)
        temp_path = temp_file.name
        temp_file.close()

        while True:
            img.save(temp_path, "JPEG", quality=quality, optimize=True)
            if os.path.getsize(temp_path) < cls.MAX_IMAGE_SIZE or quality <= 20:
                break
            quality -= 10

        return temp_path

    @classmethod
    def extract_text_from_pdf_page_image(cls, page_image_path: str) -> str:
        """从 PDF 页面图片中提取文本（理科优化模式）"""
        return cls.extract_text_from_image(page_image_path, preprocess=True, strategy="science")


class DocumentParser:
    """
    文档解析器，支持 PDF、TXT、MD、图片等格式

    双轨解析架构：
    - 轨道A: PyMuPDF4LLM（矢量PDF，保留表格/图片结构）
    - 轨道B: Vision OCR（扫描件/照片，接力式多图处理）
    """

    MIN_TEXT_THRESHOLD = 50

    # ==================== Track A: 矢量PDF解析 ====================

    @classmethod
    def parse_pdf_vectorized(cls, file_path: str, doc_id: str = None) -> Tuple[str, List[dict], dict]:
        """
        轨道A: 使用 PyMuPDF4LLM 解析矢量PDF
        保留表格结构和图片引用

        对于文本层缺失/稀少的页面（扫描件/图片页），自动触发 Vision OCR 补充。

        Returns:
            (markdown_text, page_mappings, assets_map)
        """
        if doc_id is None:
            doc_id = uuid.uuid4().hex[:12]

        try:
            import pymupdf4llm
        except ImportError:
            logger.info("pymupdf4llm 未安装，回退到 pdfplumber...")
            md_text, page_mappings = cls._parse_pdf_pdfplumber(file_path, doc_id)
            return md_text, page_mappings, {}

        # 使用 pymupdf4llm 转换为 Markdown（保留表格语法）
        markdown_text = pymupdf4llm.to_markdown(file_path)

        # 提取图片
        all_images = []
        try:
            md_with_images, extracted_images = extract_images_from_pdf(file_path, doc_id)
            if extracted_images:
                save_image_records(doc_id, extracted_images)
                image_refs = get_image_record_markdown(doc_id)
                if image_refs:
                    markdown_text = markdown_text + "\n\n" + image_refs
                all_images = extracted_images
                logger.info(f"从 PDF 提取了 {len(extracted_images)} 张图片")
        except Exception as img_err:
            logger.info(f"PDF 图片提取失败: {img_err}")

        # === 检测并 OCR 图像-only 页面 ===
        ocr_text_parts = cls._ocr_image_only_pages(file_path, markdown_text)
        if ocr_text_parts:
            ocr_append = "\n\n".join(ocr_text_parts)
            markdown_text = markdown_text + "\n\n## OCR 补充内容\n\n" + ocr_append
            logger.info(f"Vision OCR 补充了 {len(ocr_text_parts)} 个页面的内容")

        # 构建 page_mappings
        page_mappings = []
        current_pos = 0
        pages = markdown_text.split("\n\n")
        for i, page_content in enumerate(pages):
            page_mappings.append({
                "page_num": i + 1,
                "content": page_content,
                "start_char": current_pos,
                "end_char": current_pos + len(page_content)
            })
            current_pos += len(page_content) + 2

        # 构建 assets_map
        assets_map = {}
        for img_info in all_images:
            assets_map[img_info.image_id] = {
                "asset_id": img_info.image_id,
                "asset_type": "image",
                "file_path": img_info.stored_path,
                "mime_type": "image/png",
                "page_num": None,
                "sequence": len(assets_map)
            }

        # 检测是否需要 fallback
        if len(markdown_text.strip()) == 0:
            raise FallbackToOCRError("PyMuPDF4LLM 返回空结果，触发 OCR 回退")

        return markdown_text, page_mappings, assets_map

    @classmethod
    def _parse_pdf_pdfplumber(cls, file_path: str, doc_id: str = None) -> Tuple[str, List[dict]]:
        """使用 pdfplumber 解析 PDF（兼容旧方法）"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber 未安装，请运行: pip install pdfplumber")

        if doc_id is None:
            doc_id = uuid.uuid4().hex[:12]

        page_mappings = []
        all_text = []
        current_char_pos = 0

        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    page_text = page.extract_text() or ""

                    if len(page_text.strip()) < cls.MIN_TEXT_THRESHOLD:
                        logger.info(f"第 {page_num} 页文本量异常少，启用 Vision OCR...")
                        page_text = cls._parse_pdf_page_with_vision(file_path, page_num)
                    else:
                        page_text = cls._convert_to_markdown(page_text, page_num)

                    start_char = current_char_pos
                    end_char = current_char_pos + len(page_text)

                    page_mappings.append({
                        "page_num": page_num,
                        "content": page_text,
                        "start_char": start_char,
                        "end_char": end_char
                    })

                    all_text.append(page_text)
                    current_char_pos = end_char

            try:
                import fitz
                md_with_images, extracted_images = extract_images_from_pdf(file_path, doc_id)
                if extracted_images:
                    save_image_records(doc_id, extracted_images)
                    image_refs = get_image_record_markdown(doc_id)
                    all_text.append(image_refs)
                    logger.info(f"从 PDF 提取了 {len(extracted_images)} 张图片")
            except Exception as img_err:
                logger.info(f"PDF 图片提取失败: {img_err}")

        except Exception as e:
            logger.info(f"PDF 解析失败: {e}")
            all_text, page_mappings = cls._parse_pdf_with_vision_fallback(file_path)

        return "\n\n".join(all_text), page_mappings

    @classmethod
    def parse_pdf(cls, file_path: str, doc_id: str = None) -> Tuple[str, List[dict]]:
        """
        从 PDF 文件中提取文本，返回 (markdown_content, page_mappings)
        优先尝试轨道A（PyMuPDF4LLM），失败后回退到 pdfplumber
        """
        try:
            md_text, page_mappings, assets_map = cls.parse_pdf_vectorized(file_path, doc_id)
            return md_text, page_mappings
        except FallbackToOCRError:
            logger.info("轨道A结果为空，回退到 pdfplumber 解析...")
            return cls._parse_pdf_pdfplumber(file_path, doc_id)
        except Exception as e:
            logger.info(f"轨道A解析失败: {e}，回退到 pdfplumber...")
            return cls._parse_pdf_pdfplumber(file_path, doc_id)

    @classmethod
    async def parse_image_sequence_relay(
        cls, image_paths: List[str], doc_id: str = None
    ) -> Tuple[str, List[dict], dict]:
        """
        轨道B: 多图接力处理
        每张图继承前一张的摘要，保持上下文连贯
        """
        if doc_id is None:
            doc_id = uuid.uuid4().hex[:12]

        accumulated_text = []
        page_mappings = []
        assets_map = {}
        prev_summary = ""
        current_pos = 0

        for i, img_path in enumerate(image_paths):
            slices = VisionOCR.slice_long_image(img_path)
            needs_slicing = len(slices) > 1

            for j, slice_path in enumerate(slices):
                is_continuation = (i > 0 or j > 0)

                ocr_result = await VisionOCR.extract_with_context(
                    slice_path,
                    context=prev_summary,
                    is_continuation=is_continuation
                )

                if needs_slicing and slice_path != img_path:
                    try:
                        os.remove(slice_path)
                    except:
                        pass

                page_num = i + 1
                md_text = cls._convert_to_markdown(ocr_result, page_num)
                accumulated_text.append(md_text)

                page_mappings.append({
                    "page_num": page_num,
                    "content": md_text,
                    "start_char": current_pos,
                    "end_char": current_pos + len(md_text)
                })
                current_pos += len(md_text) + 2

                prev_summary = cls._generate_page_summary(ocr_result)

        full_markdown = "\n\n".join(accumulated_text)
        return full_markdown, page_mappings, assets_map

    @classmethod
    def _generate_page_summary(cls, text: str) -> str:
        """提取当前页的关键信息作为接力摘要"""
        lines = text.strip().split("\n")
        titles = [l.strip("# ") for l in lines if l.startswith("#")]
        formula_refs = re.findall(r'\(公式\s*[\d.]+\)', text)

        summary_parts = []
        if titles:
            summary_parts.append(f"标题: {'; '.join(titles[:3])}")
        if formula_refs:
            summary_parts.append(f"公式: {', '.join(formula_refs[:3])}")
        if text.strip():
            summary_parts.append(f"内容: {text.strip()[:200]}...")

        return "\n".join(summary_parts) if summary_parts else text.strip()[:300]

    # ==================== 辅助方法 ====================

    @classmethod
    def _convert_to_markdown(cls, text: str, page_num: int) -> str:
        """将提取的文本转换为 Markdown 格式，保留结构"""
        lines = text.split('\n')
        markdown_lines = [f"# 第 {page_num} 页\n"]

        for line in lines:
            line = line.strip()
            if not line:
                markdown_lines.append("")
                continue

            if re.match(r'^(第[一二三四五六七八九十\d]+[章节条项]|[①②③④⑤⑥⑦⑧⑨⑩])\s*', line):
                heading_match = re.search(r'([^\s].+)$', line)
                if heading_match:
                    markdown_lines.append(f"## {heading_match.group(1)}")
                else:
                    markdown_lines.append(f"## {line}")
            elif re.match(r'^[\d]+[.、]\s+', line) or re.match(r'^[a-zA-Z][.、]\s+', line):
                content = re.sub(r'^[\d]+[.、]\s+', '', line)
                content = re.sub(r'^[a-zA-Z][.、]\s+', '', content)
                markdown_lines.append(f"1. {content}")
            elif '|' in line and line.count('|') >= 2:
                markdown_lines.append(line)
            else:
                markdown_lines.append(line)

        return "\n".join(markdown_lines)

    @classmethod
    def _parse_pdf_page_with_vision(cls, file_path: str, page_num: int) -> str:
        """使用 Vision OCR 处理单个 PDF 页面"""
        try:
            import fitz
        except ImportError:
            raise ImportError("PyMuPDF 未安装，请运行: pip install PyMuPDF")

        try:
            pdf_document = fitz.open(file_path)
            if page_num > len(pdf_document):
                return ""

            page = pdf_document[page_num - 1]
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)

            temp_image = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            temp_path = temp_image.name
            temp_image.close()
            pix.save(temp_path)
            pdf_document.close()

            try:
                page_text = VisionOCR.extract_text_from_pdf_page_image(temp_path)
                return cls._convert_to_markdown(page_text, page_num)
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        except Exception as e:
            logger.info(f"第 {page_num} 页 Vision OCR 失败: {e}")
            return f"[第 {page_num} 页 OCR 提取失败]"

    @classmethod
    def _ocr_image_only_pages(cls, file_path: str, markdown_text: str) -> List[str]:
        """
        检测 PDF 中文本层缺失/稀少的页面，使用 Vision OCR 补充。

        对每一页检查 fitz.Page.get_text() 返回的字符数，
        如果低于阈值，则渲染为图片并调用 Vision OCR。

        Args:
            file_path: PDF 文件路径
            markdown_text: pymupdf4llm 已提取的 Markdown 文本（用于对照）

        Returns:
            OCR 补充文本列表，每项格式为 "### 第 N 页\n{ocr_text}"
        """
        try:
            import fitz
        except ImportError:
            return []

        ocr_results = []
        temp_paths = []

        try:
            pdf_document = fitz.open(file_path)

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                page_text = page.get_text() or ""

                # 如果页面文本量低于阈值，判定为图片页，需要 OCR
                if len(page_text.strip()) < cls.MIN_TEXT_THRESHOLD:
                    logger.info(f"第 {page_num + 1} 页文本量不足 ({len(page_text.strip())} 字符)，启用 Vision OCR...")

                    # 渲染为图片
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)

                    temp_image = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                    temp_path = temp_image.name
                    temp_image.close()
                    pix.save(temp_path)
                    temp_paths.append(temp_path)

                    # 调用 Vision OCR
                    try:
                        ocr_text = VisionOCR.extract_text_from_pdf_page_image(temp_path)
                        if ocr_text.strip():
                            ocr_results.append(f"### 第 {page_num + 1} 页\n\n{ocr_text.strip()}")
                            logger.info(f"第 {page_num + 1} 页 OCR 完成，提取 {len(ocr_text.strip())} 字符")
                        else:
                            logger.info(f"第 {page_num + 1} 页 OCR 返回空结果")
                    except Exception as ocr_err:
                        logger.warning(f"第 {page_num + 1} 页 Vision OCR 失败: {ocr_err}")

            pdf_document.close()

        finally:
            for temp_path in temp_paths:
                try:
                    os.remove(temp_path)
                except:
                    pass

        return ocr_results

    @classmethod
    def _parse_pdf_with_vision_fallback(cls, file_path: str) -> Tuple[str, List[dict]]:
        """使用 Vision OCR 处理整个 PDF（回退方案）"""
        try:
            import fitz
        except ImportError:
            raise ImportError("PyMuPDF 未安装，请运行: pip install PyMuPDF")

        page_mappings = []
        all_text = []
        temp_image_paths = []

        try:
            pdf_document = fitz.open(file_path)

            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)

                temp_image = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                temp_path = temp_image.name
                temp_image.close()
                pix.save(temp_path)
                temp_image_paths.append(temp_path)

                page_text = VisionOCR.extract_text_from_pdf_page_image(temp_path)
                md_text = cls._convert_to_markdown(page_text, page_num + 1)

                page_mappings.append({
                    "page_num": page_num + 1,
                    "content": md_text,
                    "start_char": sum(len(t) + 2 for t in all_text),
                    "end_char": sum(len(t) + 2 for t in all_text) + len(md_text)
                })

                all_text.append(md_text)

            pdf_document.close()

        finally:
            for temp_path in temp_image_paths:
                try:
                    os.remove(temp_path)
                except:
                    pass

        return "\n\n".join(all_text), page_mappings

    @classmethod
    def parse_image(cls, file_path: str) -> Tuple[str, List[dict]]:
        """直接解析图片文件，使用 Vision LLM OCR"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"图片文件不存在: {file_path}")

        page_num = 1
        text = VisionOCR.extract_text_from_image(file_path)
        md_text = cls._convert_to_markdown(text, page_num)

        mappings = [{
            "page_num": page_num,
            "content": md_text,
            "start_char": 0,
            "end_char": len(md_text)
        }]

        return md_text, mappings

    @classmethod
    def parse_txt(cls, file_path: str) -> Tuple[str, List[dict]]:
        """从 TXT 文件读取文本"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()

        lines = text.split('\n')
        md_lines = []
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                md_lines.append("")
                continue

            if len(line) < 50 and not line[-1] in '。！？；：' and i < len(lines) - 1:
                next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
                if len(next_line) > 20:
                    md_lines.append(f"## {line}")
                    continue

            md_lines.append(line)

        md_text = "\n".join(md_lines)

        mappings = [{
            "page_num": 1,
            "content": md_text,
            "start_char": 0,
            "end_char": len(md_text)
        }]

        return md_text, mappings

    @classmethod
    def parse_docx(cls, file_path: str, doc_id: str = None) -> Tuple[str, List[dict]]:
        """从 Word .docx 文件提取文本和图片（图片插入到正确位置）"""
        from docx import Document
        import zipfile

        if doc_id is None:
            doc_id = uuid.uuid4().hex[:12]

        doc = Document(file_path)
        handler = ImageHandler(doc_id)

        # Step 1: 从 ZIP 提取所有图片并保存，建立 rId -> ImageInfo 映射
        rId_to_image = {}
        try:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                media_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
                for mf in media_files:
                    fname = os.path.basename(mf)
                    if fname:
                        try:
                            image_data = zip_ref.read(mf)
                            ext = os.path.splitext(fname)[1].lstrip('.')
                            img_info = handler.save_image(
                                image_data=image_data,
                                original_name=fname,
                                image_format=ext or 'png'
                            )
                            # 用文件名和完整路径都建立映射
                            rId_to_image[fname] = img_info
                        except Exception as e:
                            logger.info(f"提取图片失败 {mf}: {e}")
        except Exception as e:
            logger.info(f"无法打开 DOCX ZIP: {e}")

        # Step 2: 建立 rId -> 图片文件名 映射（从文档关系）
        rId_to_filename = {}
        for r_id, rel in doc.part.rels.items():
            if 'image' in str(rel.reltype).lower():
                target = rel.target_ref
                rId_to_filename[r_id] = os.path.basename(target)

        # Step 3: 为每个段落收集其图片引用
        para_image_refs = {}  # para_index -> [img_md, ...]
        for shape in doc.inline_shapes:
            try:
                blip = shape._inline.graphic.graphicData.pic.blipFill.blip
                r_id = blip.embed
                fname = rId_to_filename.get(r_id)
                if fname and fname in rId_to_image:
                    img = rId_to_image[fname]
                    img_md = f"![{img.original_name}]({img.stored_path.replace(chr(92), '/')})"
                    # 从 shape._inline 向上找到 w:p 段落元素
                    drawing_el = shape._inline.getparent()
                    current = drawing_el
                    para_el = None
                    while current is not None:
                        if current.tag.endswith('}p'):
                            para_el = current
                            break
                        current = current.getparent()
                    # 找到对应的段落索引
                    if para_el is not None:
                        for p_idx, para in enumerate(doc.paragraphs):
                            if para._element is para_el:
                                para_image_refs.setdefault(p_idx, []).append(img_md)
                                break
            except Exception:
                pass

        # Step 4: 遍历段落，插入文本和图片
        paragraphs = []
        mappings = []
        current_pos = 0

        for para_idx, para in enumerate(doc.paragraphs):
            # 先插入该段落的图片
            for img_md in para_image_refs.get(para_idx, []):
                paragraphs.append(img_md)

            text = para.text.strip()
            if text:
                # 检测标题样式
                is_heading = False
                heading_level = 2
                if para.style.name.startswith('Heading'):
                    is_heading = True
                    level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 2
                    heading_level = min(level, 6)
                elif re.match(r'^[一二三四五六七八九十]+[、.]', text):
                    is_heading = True
                    heading_level = 1
                elif re.match(r'^\d+[、.]\s', text):
                    is_heading = True
                    heading_level = 2
                elif re.match(r'^\d+\.\d+[、.]\s', text):
                    is_heading = True
                    heading_level = 3
                elif re.match(r'^第[一二三四五六七八九十\d]+[章节条]', text):
                    is_heading = True
                    heading_level = 1

                if is_heading:
                    prefix = "#" * heading_level
                    text = f"{prefix} {text}"

                start_char = current_pos
                end_char = current_pos + len(text)

                paragraphs.append(text)
                mappings.append({
                    "page_num": len(mappings) + 1,
                    "content": text,
                    "start_char": start_char,
                    "end_char": end_char
                })
                current_pos = end_char + 2

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells])
                if row_text:
                    paragraphs.append(f"| {row_text} |")
                    mappings.append({
                        "page_num": len(mappings) + 1,
                        "content": f"| {row_text} |",
                        "start_char": current_pos,
                        "end_char": current_pos + len(row_text) + 4
                    })
                    current_pos += len(row_text) + 4

        md_text = "\n\n".join(paragraphs)

        # 记录图片
        if handler.images:
            save_image_records(doc_id, handler.images)
            logger.info(f"从 DOCX 提取了 {len(handler.images)} 张图片")

        return md_text, mappings

    @classmethod
    def parse_doc(cls, file_path: str) -> Tuple[str, List[dict]]:
        """从 Word .doc 文件提取文本（支持 Windows 和 Linux）"""
        import platform
        system = platform.system()

        # Windows: 使用 win32com
        if system == "Windows":
            try:
                import win32com.client
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    text = doc.Content.Text
                    doc.Close(False)
                    word.Quit()
                    return cls._text_to_markdown(text)
                finally:
                    pythoncom.CoUninitialize()
            except ImportError:
                raise ImportError("pywin32 未安装，请运行: pip install pywin32")
            except Exception as e:
                raise Exception(f"无法解析 .doc 文件: {str(e)}")

        # Linux: 使用 antiword
        else:
            import shutil
            import subprocess

            if not shutil.which("antiword"):
                raise Exception("antiword 未安装，请在 Railway Build Command 中添加: apt-get install -y antiword")

            try:
                result = subprocess.run(
                    ["antiword", "-w", "0", os.path.abspath(file_path)],
                    capture_output=True, text=True, timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    return cls._text_to_markdown(result.stdout)
                else:
                    raise Exception(f"antiword 解析失败，返回码: {result.returncode}")
            except subprocess.TimeoutExpired:
                raise Exception("antiword 解析超时")
            except Exception as e:
                raise Exception(f"无法解析 .doc 文件: {str(e)}")

    @classmethod
    def _text_to_markdown(cls, text: str) -> Tuple[str, List[dict]]:
        """将纯文本转换为带结构的 Markdown"""
        lines = text.split('\n')
        md_lines = []
        mappings = []
        current_pos = 0

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                current_pos += 1
                md_lines.append("")
                continue

            if re.match(r'^(第[一二三四五六七八九十\d]+[章节条项部]|第[一二三四五六七八九十\d]+讲)\s+', line):
                md_lines.append(f"## {line}")
            elif re.match(r'^\d+\.\s+[^\d]', line):
                cleaned = re.sub(r'^\d+\.\s+', '', line)
                md_lines.append(f"## {cleaned}")
            else:
                md_lines.append(line)

            start = current_pos
            end = current_pos + len(line)
            mappings.append({
                "page_num": i // 30 + 1,
                "content": line,
                "start_char": start,
                "end_char": end
            })
            current_pos = end + 2

        return "\n".join(md_lines), mappings

    @classmethod
    def parse_xlsx(cls, file_path: str) -> Tuple[str, List[dict]]:
        """从 Excel .xlsx 文件提取文本"""
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        result_parts = []
        mappings = []
        current_pos = 0

        for sheet_idx, sheet_name in enumerate(wb.sheetnames, start=1):
            sheet = wb[sheet_name]
            result_parts.append(f"## {sheet_name}\n")
            start = current_pos
            current_pos += len(f"## {sheet_name}\n")

            rows_data = []
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v for v in row_values):
                    row_text = " | ".join(row_values) + " |"
                    rows_data.append(f"| {row_text} |")

                    mappings.append({
                        "page_num": sheet_idx,
                        "content": row_text,
                        "start_char": current_pos,
                        "end_char": current_pos + len(row_text)
                    })
                    current_pos += len(row_text) + 1

            if rows_data:
                col_count = len(rows_data[0].split(" | ")) if " | " in rows_data[0] else 1
                result_parts.append("| " + " | ".join(["---"] * col_count) + " |\n")
                result_parts.extend([r + "\n" for r in rows_data])
                result_parts.append("\n")

        wb.close()
        md_text = "".join(result_parts)

        return md_text, mappings

    @classmethod
    def parse_xls(cls, file_path: str) -> Tuple[str, List[dict]]:
        """从 Excel .xls 文件提取文本"""
        import xlrd

        wb = xlrd.open_workbook(file_path)
        result_parts = []
        mappings = []
        current_pos = 0

        for sheet_idx in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_idx)
            result_parts.append(f"## {sheet.name}\n")
            start = current_pos
            current_pos += len(f"## {sheet.name}\n")

            for row_idx in range(sheet.nrows):
                row_values = []
                for col_idx in range(sheet.ncols):
                    cell = sheet.cell(row_idx, col_idx)
                    row_values.append(str(cell.value) if cell.value else "")
                if any(v for v in row_values):
                    row_text = " | ".join(row_values) + " |"
                    mappings.append({
                        "page_num": sheet_idx + 1,
                        "content": row_text,
                        "start_char": current_pos,
                        "end_char": current_pos + len(row_text)
                    })
                    current_pos += len(row_text) + 1

        md_text = "".join(result_parts)
        return md_text, mappings

    @classmethod
    def parse_file(cls, file_path: str, file_extension: str, doc_id: str = None) -> Tuple[str, List[dict]]:
        """
        根据文件扩展名调用对应解析器

        双轨分发：
        - PDF: 优先轨道A（PyMuPDF4LLM），自动 fallback 到 pdfplumber/Vision OCR
        - 图片: 轨道B（Vision OCR）

        Returns:
            tuple: (markdown格式文本, page_mappings)
        """
        if doc_id is None:
            doc_id = uuid.uuid4().hex[:12]

        ext = file_extension.lower()
        if ext == ".pdf":
            return cls.parse_pdf(file_path, doc_id)
        elif ext in [".txt", ".md"]:
            return cls.parse_txt(file_path)
        elif ext == ".docx":
            return cls.parse_docx(file_path, doc_id)
        elif ext == ".doc":
            return cls.parse_doc(file_path)
        elif ext == ".xlsx":
            return cls.parse_xlsx(file_path)
        elif ext == ".xls":
            return cls.parse_xls(file_path)
        elif ext in [".png", ".jpg", ".jpeg"]:
            return cls.parse_image(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")


class RAGEngine:
    """
    RAG 引擎核心类
    使用 LangChain + Chroma 实现多表示索引 + BM25 混合搜索
    """

    def __init__(self):
        self.embeddings = self._init_embeddings()
        self.vectorstore = self._init_vectorstore()
        self.id_key = "doc_id"
        self.retriever = self._init_retriever()
        self.bm25_retriever = self._init_bm25_retriever()
        self.reranker = self._init_reranker()

    def _init_reranker(self) -> RerankerClient:
        """初始化 Reranker 客户端"""
        return RerankerClient(reranker_url=settings.RERANKER_URL)

    def _init_bm25_retriever(self) -> BM25Retriever:
        """初始化 BM25 检索器"""
        bm25_retriever = BM25Retriever(persist_path=settings.BM25_PERSIST_PATH)
        bm25_retriever.load()
        return bm25_retriever

    def _init_embeddings(self):
        """初始化 Embedding 模型"""
        return DashScopeEmbeddings(
            model=settings.EMBEDDING_MODEL,
            api_key=settings.OPENAI_API_KEY if settings.OPENAI_API_KEY else "sk-mock",
            base_url=settings.OPENAI_API_BASE
        )

    def _init_vectorstore(self) -> Chroma:
        """初始化 Chroma 向量数据库"""
        os.makedirs(settings.CHROMA_DB_PATH, exist_ok=True)
        return Chroma(
            persist_directory=settings.CHROMA_DB_PATH,
            embedding_function=self.embeddings,
            collection_name="document_chunks"
        )

    def _init_retriever(self) -> MultiVectorRetriever:
        """初始化多向量检索器"""
        vectorstore = self.vectorstore
        docstore = InMemoryStore()
        retriever = MultiVectorRetriever(
            vectorstore=vectorstore,
            docstore=docstore,
            id_key=self.id_key,
            search_type="similarity",
            search_kwargs={"k": 5}
        )
        return retriever

    def process_document(
        self,
        file_path: str,
        file_extension: str,
        metadata: Optional[dict] = None,
        doc_id: Optional[str] = None
    ) -> dict:
        """
        处理文档：解析 -> 切分 -> 入库（Multi-vector Retriever）

        实现 Parent-Child 结构：
        - 母文档（完整上下文）存入 docstore
        - 子文档（检索用）存入向量数据库

        Args:
            doc_id: 可选，外部传入的文档ID（如未提供则自动生成）

        Returns:
            dict: {
                "markdown": markdown格式文本,
                "page_mappings": [{page_num, content, start_char, end_char}, ...],
                "doc_id": 文档ID,
                "status": "OK" | "WARNING",
                "verification_warnings": [...]
            }
        """
        doc_id = doc_id or str(uuid.uuid4())

        # 解析文档
        markdown_text, page_mappings = DocumentParser.parse_file(file_path, file_extension, doc_id)

        # 获取图片元数据
        from core.image_handler import get_image_records, get_image_record_markdown
        image_records = get_image_records(doc_id)

        # 追加图片引用（仅当内容中尚未包含时，如 DOCX 已内联插入）
        if image_records:
            image_refs = get_image_record_markdown(doc_id)
            if image_refs and image_refs.strip() not in markdown_text:
                # 检查是否已有图片引用内联在内容中
                existing_img_refs = re.findall(r'!\[.*?\]\(.*?\)', markdown_text)
                if not existing_img_refs:
                    markdown_text = markdown_text + "\n\n" + image_refs

        # 注入 RAG 语义增强标签到表格
        table_summaries = self._inject_table_rag_tags(markdown_text)

        # 抽取多种表示（结构、知识点等）供文档概览使用
        try:
            from core.doc_bot import get_doc_bot_v2
            doc_bot = get_doc_bot_v2()
            representations = doc_bot.extractor.extract(markdown_text, doc_id, page_mappings)
            doc_bot.vector_store.add_representations(representations)
        except Exception as e:
            logger.info(f"表示抽取失败: {e}")
            import traceback
            traceback.print_exc()

        # 切分文档并入库
        self._split_and_index(markdown_text, doc_id, metadata, table_summaries)

        # 自校对
        warnings = self._self_verify(markdown_text, image_records)

        result = {
            "doc_id": doc_id,
            "markdown": markdown_text,
            "page_mappings": page_mappings,
            "raw_text_length": len(markdown_text),
            "image_count": len(image_records)
        }

        if warnings:
            result["status"] = "WARNING"
            result["verification_warnings"] = warnings

        return result

    def _inject_table_rag_tags(self, markdown_text: str) -> List[dict]:
        """
        Patch 4: RAG语义增强标签
        检测 Markdown 中的表格块，注入 HTML 注释形式的语义标签
        """
        table_summaries = []

        # 查找所有 Markdown 表格块
        table_pattern = re.compile(
            r'((?:^\|.*\|$\n)+)',
            re.MULTILINE
        )

        for match in table_pattern.finditer(markdown_text):
            table_block = match.group(1)
            lines = table_block.strip().split('\n')

            if len(lines) < 2:
                continue

            # 提取表头
            headers = [h.strip() for h in lines[0].strip('|').split('|')]
            # 估算行数（减去分隔线行）
            row_count = max(0, len(lines) - 2)

            # 提取首行和末行数据用于趋势分析
            data_lines = [l for l in lines if not re.match(r'^\|[\s\-:|]+\|$', l.strip()) and '|' in l]
            first_row = []
            last_row = []
            if len(data_lines) > 0:
                first_row = [c.strip() for c in data_lines[0].strip('|').split('|')]
            if len(data_lines) > 1:
                last_row = [c.strip() for c in data_lines[-1].strip('|').split('|')]

            # 生成趋势描述
            trend = ""
            if first_row and last_row and len(first_row) > 1 and len(last_row) > 1:
                try:
                    first_num = float(first_row[-1].replace(',', '').replace('%', ''))
                    last_num = float(last_row[-1].replace(',', '').replace('%', ''))
                    if last_num > first_num:
                        trend = f"增长趋势: {first_row[0]}→{last_row[0]} 末值上升"
                    elif last_num < first_num:
                        trend = f"下降趋势: {first_row[0]}→{last_row[0]} 末值下降"
                except (ValueError, IndexError):
                    trend = ""

            # 生成 HTML 版本
            html_lines = ["<table>"]
            html_lines.append("  <tr>" + "".join(f"<th>{h}</th>" for h in headers) + "</tr>")
            for line in data_lines[1:] if len(data_lines) > 1 else data_lines:
                cells = [c.strip() for c in line.strip('|').split('|')]
                html_lines.append("  <tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
            html_lines.append("</table>")
            table_html = "\n".join(html_lines)

            # 生成结构化 JSON
            table_json = json.dumps({
                "headers": headers,
                "rows": [
                    [c.strip() for c in l.strip('|').split('|')]
                    for l in data_lines
                ],
                "row_count": row_count,
                "context": trend,
                "format": "markdown_table"
            }, ensure_ascii=False)

            # 生成 RAG 语义增强标签（HTML 注释，不影响渲染）
            topic = headers[0] if headers else "未知"
            rag_tag = (
                f"<!-- RAG_TABLE_SUMMARY: 主题=\"{topic}\"; "
                f"表头=[{', '.join(headers)}]; "
                f"趋势=\"{trend}\"; "
                f"行数={row_count} -->"
            )

            # 记录表格信息
            table_summaries.append({
                "position": match.start(),
                "table_html": table_html,
                "table_json": table_json,
                "rag_tag": rag_tag
            })

        # 从后向前注入标签（避免位置偏移）
        for table_info in reversed(table_summaries):
            pos = table_info["position"]
            markdown_text = markdown_text[:pos] + table_info["rag_tag"] + "\n" + markdown_text[pos:]

        return table_summaries

    def _split_and_index(self, markdown_text: str, doc_id: str, metadata: Optional[dict] = None, table_summaries: Optional[List[dict]] = None) -> None:
        """
        将文档切分为母文档和子文档，并分别入库

        - 母文档（Parent）：完整上下文，存入 docstore
        - 子文档（Child）：检索用小块，存入向量数据库
        """
        parent_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2000,
            chunk_overlap=100,
            separators=["\n\n", "\n", "## ", "### ", "# ", "."]
        )
        child_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400,
            chunk_overlap=50,
            separators=["\n\n", "\n", "## ", "### ", "# "]
        )

        parent_docs = parent_splitter.create_documents([markdown_text])

        docstore_kv_pairs = []
        child_docs_to_add = []

        # 构建表格 metadata 索引（用于快速查找包含表格的母文档）
        table_metadata = {}
        if table_summaries:
            table_metadata["table_count"] = len(table_summaries)
            table_metadata["table_html"] = "\n\n".join(t["table_html"] for t in table_summaries)
            table_metadata["table_json"] = "\n\n".join(t["table_json"] for t in table_summaries)
            table_metadata["table_summary"] = "\n".join(t["rag_tag"] for t in table_summaries)

        for parent_doc in parent_docs:
            parent_id = str(uuid.uuid4())

            parent_doc.metadata["parent_id"] = parent_id
            parent_doc.metadata["doc_id"] = doc_id
            if metadata:
                parent_doc.metadata.update(metadata)
            # 注入表格 metadata
            if table_metadata:
                parent_doc.metadata.update(table_metadata)

            docstore_kv_pairs.append((parent_id, parent_doc))

            child_docs = child_splitter.split_documents([parent_doc])

            for child_doc in child_docs:
                # 过滤空文本块，防止 DashScope 400 错误
                if not child_doc.page_content or not child_doc.page_content.strip():
                    continue
                child_doc.metadata["parent_id"] = parent_id
                child_doc.metadata["doc_id"] = doc_id
                if metadata:
                    child_doc.metadata.update(metadata)
                if table_metadata:
                    child_doc.metadata.update(table_metadata)
                child_docs_to_add.append(child_doc)

        if child_docs_to_add:
            self.vectorstore.add_documents(child_docs_to_add)
            logger.info(f"[EMBED] Text RAG] Added {len(child_docs_to_add)} child docs to vectorstore")

            # 同步构建 BM25 索引
            if settings.BM25_ENABLED:
                self.bm25_retriever.add_documents(child_docs_to_add)
                self.bm25_retriever.save()
                logger.info(f"[BM25] 已将 {len(child_docs_to_add)} 个子文档加入 BM25 索引")

        if docstore_kv_pairs:
            self.retriever.docstore.mset(docstore_kv_pairs)
            logger.info(f"[EMBED] Text RAG] Added {len(docstore_kv_pairs)} parent docs to docstore")

    def _self_verify(self, markdown: str, image_records: list) -> List[str]:
        """
        Patch 5: Asset 数量自校对
        对比提取的 assets 与 Markdown 中的占位符数量
        """
        warnings = []

        # 统计图片占位符
        md_image_refs = re.findall(r'!\[.*?\]\((.*?)\)', markdown)
        expected_images = len(image_records)

        if expected_images > 0 and len(md_image_refs) != expected_images:
            missing_ids = [img.image_id for img in image_records
                         if not any(img.stored_path in ref for ref in md_image_refs)]
            warnings.append(
                f"IMAGE_MISMATCH: 提取了{expected_images}张图片，"
                f"但Markdown中只有{len(md_image_refs)}个占位符。"
                f"缺失: {missing_ids}"
            )

        # 统计表格
        md_tables = re.findall(r'\|.*?\|.*?\n\|[-:|]+\|', markdown)
        expected_tables = len([r for r in image_records if hasattr(r, 'asset_type') and getattr(r, 'asset_type', None) == 'table'])

        if expected_tables > 0 and len(md_tables) != expected_tables:
            warnings.append(
                f"TABLE_MISMATCH: 提取了{expected_tables}个表格，"
                f"但Markdown中只找到{len(md_tables)}个"
            )

        return warnings

    def similarity_search(self, query: str, k: int = 5) -> list[Document]:
        """执行相似度搜索"""
        return self.vectorstore.similarity_search(query, k=k)

    def retrieve_with_expansion(self, query: str, k: int = 3) -> list[Document]:
        """检索并扩展上下文"""
        sub_docs = self.vectorstore.similarity_search(query, k=k)

        parent_ids = set()
        for doc in sub_docs:
            doc_id = doc.metadata.get("doc_id")
            if doc_id:
                parent_ids.add(doc_id)

        all_docs = self.vectorstore.get()
        parent_docs = []

        for doc in all_docs["documents"]:
            meta = all_docs["metadatas"][all_docs["documents"].index(doc)]
            if meta.get("doc_id") in parent_ids and "parent_id" not in meta:
                parent_docs.append(Document(page_content=doc, metadata=meta))

        if not parent_docs:
            return sub_docs

        return parent_docs

    def get_document_by_id(self, doc_id: str) -> list[Document]:
        """根据 doc_id 获取文档的所有块"""
        return self.vectorstore.get(where={"doc_id": doc_id})

    def _rrf_fusion(self, vector_results: list[Document], bm25_results: list[Document], k: int = 60) -> list[Document]:
        """
        Reciprocal Rank Fusion (RRF): 融合向量检索和 BM25 检索结果
        score = Σ 1/(k + rank)
        """
        from collections import defaultdict

        doc_map: dict[str, Document] = {}
        for doc in vector_results:
            doc_map[doc.page_content] = doc
        for doc in bm25_results:
            doc_map[doc.page_content] = doc

        scores: dict[str, float] = defaultdict(float)

        for rank, doc in enumerate(vector_results):
            scores[doc.page_content] += 1.0 / (k + rank + 1)

        for rank, doc in enumerate(bm25_results):
            scores[doc.page_content] += 1.0 / (k + rank + 1)

        sorted_contents = sorted(scores.items(), key=lambda x: x[1], reverse=True)
        return [doc_map[content] for content, _ in sorted_contents]

    def hybrid_search(self, query: str, k: int = 5, mode: str = "hybrid", use_rerank: bool = True) -> list[Document]:
        """
        混合搜索入口

        Args:
            query: 查询文本
            k: 返回结果数量
            mode: 'vector' | 'bm25' | 'hybrid' (默认 hybrid)
                - vector: 仅向量检索
                - bm25: 仅 BM25 关键词检索
                - hybrid: RRF 融合向量 + BM25
            use_rerank: 是否使用 Reranker 二阶段重排（默认 True）
        """
        if mode == "vector":
            results = self.vectorstore.similarity_search(query, k=k)
        elif mode == "bm25":
            results = self.bm25_retriever.search(query, k=k)
        else:  # hybrid
            # 扩展检索数量，避免融合后结果太少
            expand_k = max(k * 3, 20)
            vector_results = self.vectorstore.similarity_search(query, k=expand_k)
            bm25_results = self.bm25_retriever.search(query, k=expand_k)

            if not vector_results and not bm25_results:
                return []
            if not vector_results:
                results = bm25_results
            elif not bm25_results:
                results = vector_results
            else:
                fused = self._rrf_fusion(vector_results, bm25_results, k=60)
                # RRF 融合后取 2*k 条送 reranker，留足够候选
                results = fused[:k * 2]

        # 第二阶段：Reranker 精排
        if use_rerank and settings.RERANKER_ENABLED and results:
            results = self.reranker.rerank(query, results, top_n=k)

        return results[:k]


# 全局 RAG 引擎实例
_rag_engine: Optional[RAGEngine] = None


def get_rag_engine() -> RAGEngine:
    """获取 RAG 引擎单例"""
    global _rag_engine
    if _rag_engine is None:
        _rag_engine = RAGEngine()
    return _rag_engine

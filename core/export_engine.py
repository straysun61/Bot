"""
MD 导出引擎 - 将 Markdown 转换为其他格式
独立模块，不影响原 MD 转换和 RAG 功能
"""
import os
import re
import uuid
import logging
from typing import Optional, Dict, List
from datetime import datetime

logger = logging.getLogger(__name__)

# 存储目录
EXPORT_STORAGE_DIR = "./storage_export"

# 确保导出目录存在
os.makedirs(EXPORT_STORAGE_DIR, exist_ok=True)


class ExportFormat:
    """支持的导出格式"""
    DOCX = "docx"
    PDF = "pdf"
    HTML = "html"
    TXT = "txt"


class MarkdownExporter:
    """
    Markdown 导出器
    支持将 MD 转换为 Word、PDF、HTML、纯文本
    """

    def __init__(self):
        self.export_dir = EXPORT_STORAGE_DIR

    def _extract_images_from_markdown(self, md_content: str) -> List[Dict]:
        """从 Markdown 中提取所有图片引用"""
        pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        matches = re.findall(pattern, md_content)

        images = []
        for alt_text, image_path in matches:
            if os.path.exists(image_path):
                images.append({
                    "alt": alt_text,
                    "path": image_path,
                    "original_ref": f"![{alt_text}]({image_path})"
                })
            else:
                logger.warning(f"图片文件不存在: {image_path}")
        return images

    def _process_markdown_for_docx(self, md_content: str) -> str:
        """处理 Markdown 内容，将图片路径转换为 docx 能识别的格式"""
        # 提取图片引用
        images = self._extract_images_from_markdown(md_content)

        # 保留图片引用（docx导出时会处理）
        # 返回处理后的内容
        return md_content

    def export_to_docx(self, doc_id: str, md_content: str) -> Dict:
        """
        将 MD 转换为 Word (.docx)

        Args:
            doc_id: 文档ID
            md_content: Markdown 内容

        Returns:
            包含导出结果的字典
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import qn
            import markdown
            from lxml import etree

            # 提取图片
            images = self._extract_images_from_markdown(md_content)
            image_map = {img["path"]: img for img in images}

            # 创建 Word 文档
            doc = Document()

            # 解析 Markdown 行
            lines = md_content.split('\n')
            in_code_block = False
            code_lines = []

            for line in lines:
                line = line.rstrip()

                # 代码块处理
                if line.strip().startswith('```'):
                    if in_code_block:
                        # 结束代码块
                        para = doc.add_paragraph()
                        para.style = 'No Spacing'
                        for code_line in code_lines:
                            run = para.add_run(code_line)
                            run.font.name = 'Consolas'
                            run.font.size = Pt(9)
                        code_lines = []
                        in_code_block = False
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_lines.append(line)
                    continue

                # 标题处理
                if line.startswith('# '):
                    level = 1
                    text = line[2:].strip()
                    heading = doc.add_heading(text, level=level)
                elif line.startswith('## '):
                    level = 2
                    text = line[3:].strip()
                    heading = doc.add_heading(text, level=level)
                elif line.startswith('### '):
                    level = 3
                    text = line[4:].strip()
                    heading = doc.add_heading(text, level=level)
                elif line.startswith('#### '):
                    level = 4
                    text = line[5:].strip()
                    heading = doc.add_heading(text, level=level)
                # 引用处理
                elif line.startswith('> '):
                    text = line[2:].strip()
                    para = doc.add_paragraph(text)
                    para_format = para.paragraph_format
                    para_format.left_indent = Cm(0.5)
                    # 添加竖线样式
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(100, 100, 100)
                # 列表处理
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    text = line.strip()[2:].strip()
                    para = doc.add_paragraph(text, style='List Bullet')
                elif re.match(r'^\d+\.\s+', line.strip()):
                    match = re.match(r'^(\d+\.)\s+(.*)', line.strip())
                    if match:
                        text = match.group(2)
                        para = doc.add_paragraph(text, style='List Number')
                # 分割线
                elif line.strip() in ['---', '***', '___']:
                    doc.add_paragraph()
                # 图片处理
                elif line.strip().startswith('!['):
                    # 解析图片引用
                    match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
                    if match:
                        alt_text = match.group(1)
                        image_path = match.group(2)
                        if os.path.exists(image_path):
                            try:
                                # 添加图片
                                para = doc.add_paragraph()
                                run = para.add_run()
                                run.add_picture(image_path, width=Inches(5))
                                # 添加图片说明
                                caption = doc.add_paragraph(alt_text)
                                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in caption.runs:
                                    run.font.size = Pt(9)
                                    run.font.color.rgb = RGBColor(128, 128, 128)
                            except Exception as img_err:
                                logger.warning(f"添加图片失败: {img_err}")
                                para = doc.add_paragraph(line.strip())
                        else:
                            para = doc.add_paragraph(f"[图片不存在: {image_path}]")
                # 表格处理
                elif '|' in line and line.strip().startswith('|'):
                    # 简单表格处理 - 收集多行
                    table_lines = [line]
                    idx = lines.index(line)
                    for next_line in lines[idx+1:idx+10]:
                        if next_line.strip().startswith('|'):
                            table_lines.append(next_line)
                        else:
                            break

                    # 移除分隔行（---|---格式）
                    data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l.strip())]

                    if len(data_lines) >= 2:
                        # 解析表格
                        rows_data = []
                        for row_line in data_lines:
                            cells = [c.strip() for c in row_line.strip().strip('|').split('|')]
                            if cells and any(cells):
                                rows_data.append(cells)

                        if rows_data:
                            cols = len(rows_data[0])
                            table = doc.add_table(rows=len(rows_data), cols=cols)
                            table.style = 'Table Grid'

                            for i, row_data in enumerate(rows_data):
                                for j, cell_text in enumerate(row_data):
                                    if j < cols:
                                        cell = table.rows[i].cells[j]
                                        cell.text = cell_text

                            doc.add_paragraph()
                # 空行
                elif not line.strip():
                    doc.add_paragraph()
                # 普通段落
                else:
                    # 处理行内格式（加粗、斜体）
                    processed_line = self._process_inline_format(line)
                    para = doc.add_paragraph(processed_line)

            # 生成文件路径
            filename = f"{doc_id}_{uuid.uuid4().hex[:8]}.docx"
            filepath = os.path.join(self.export_dir, filename)
            doc.save(filepath)

            return {
                "success": True,
                "format": ExportFormat.DOCX,
                "filename": filename,
                "filepath": filepath,
                "size": os.path.getsize(filepath),
                "images_embedded": len(images)
            }

        except ImportError as e:
            logger.error(f"缺少依赖库: {e}")
            return {
                "success": False,
                "error": f"缺少依赖库: {str(e)}",
                "format": ExportFormat.DOCX
            }
        except Exception as e:
            logger.error(f"导出 Word 失败: {e}")
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "format": ExportFormat.DOCX
            }

    def _process_inline_format(self, text: str) -> str:
        """处理行内格式（保留原始文本，不做转换）"""
        # 简单处理：移除部分格式标记，保持可读性
        # python-docx 的 run 不支持直接添加加粗等格式，所以我们保留原始文本
        return text

    def export_to_pdf(self, doc_id: str, md_content: str) -> Dict:
        """
        将 MD 转换为 PDF

        Args:
            doc_id: 文档ID
            md_content: Markdown 内容

        Returns:
            包含导出结果的字典
        """
        try:
            import markdown
            import re
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib.styles import ParagraphStyle
            from reportlab.lib.enums import TA_LEFT, TA_CENTER
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
            from reportlab.lib import colors
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            # 注册中文字体
            try:
                pdfmetrics.registerFont(TTFont('SimSun', 'C:/Windows/Fonts/simsun.ttc'))
                font_name = 'SimSun'
            except:
                try:
                    pdfmetrics.registerFont(TTFont('Microsoft YaHei', 'C:/Windows/Fonts/msyh.ttc'))
                    font_name = 'Microsoft YaHei'
                except:
                    font_name = 'Helvetica'

            # 解析 Markdown
            lines = md_content.split('\n')
            story = []

            # 定义样式
            styles = {
                'h1': ParagraphStyle('h1', fontName=font_name, fontSize=18, spaceAfter=12, spaceBefore=20, textColor=colors.HexColor('#2c3e50')),
                'h2': ParagraphStyle('h2', fontName=font_name, fontSize=14, spaceAfter=10, spaceBefore=15, textColor=colors.HexColor('#34495e')),
                'h3': ParagraphStyle('h3', fontName=font_name, fontSize=12, spaceAfter=8, spaceBefore=10, textColor=colors.HexColor('#555')),
                'body': ParagraphStyle('body', fontName=font_name, fontSize=10, spaceAfter=6, leading=14),
                'code': ParagraphStyle('code', fontName='Courier', fontSize=8, spaceAfter=6, backColor=colors.HexColor('#f4f4f4'), leftIndent=10),
            }

            # 简化处理：逐行解析
            in_code_block = False
            code_content = []
            table_rows = []
            in_table = False

            for line in lines:
                # 代码块
                if line.strip().startswith('```'):
                    if in_code_block:
                        # 结束代码块
                        story.append(Paragraph('<br/>'.join(code_content), styles['code']))
                        story.append(Spacer(1, 6))
                        code_content = []
                    else:
                        in_code_block = True
                    continue

                if in_code_block:
                    code_content.append(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'))
                    continue

                # 表格处理（简化）
                if '|' in line and line.strip().startswith('|'):
                    table_rows.append(line)
                    in_table = True
                    continue
                elif in_table:
                    # 处理表格
                    if len(table_rows) > 1:
                        header = [h.strip() for h in table_rows[0].split('|')[1:-1]]
                        data = []
                        for row in table_rows[2 if len(table_rows) > 2 else 1:]:
                            cells = [c.strip() for c in row.split('|')[1:-1]]
                            if cells and any(cells):
                                data.append(cells)
                        if data:
                            t = Table([header] + data)
                            t.setStyle(TableStyle([
                                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f5f5f5')),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#ddd')),
                                ('FONTNAME', (0, 0), (-1, -1), font_name),
                                ('FONTSIZE', (0, 0), (-1, -1), 9),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                ('PADDING', (0, 0), (-1, -1), 4),
                            ]))
                            story.append(t)
                            story.append(Spacer(1, 10))
                    table_rows = []
                    in_table = False

                # 标题
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['h1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['h2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['h3']))
                # 引用
                elif line.startswith('> '):
                    story.append(Paragraph(line[2:], styles['body']))
                # 列表
                elif line.strip().startswith('- ') or line.strip().startswith('* '):
                    story.append(Paragraph(f"• {line.strip()[2:]}", styles['body']))
                elif re.match(r'^\d+\. ', line.strip()):
                    story.append(Paragraph(line.strip(), styles['body']))
                # 分割线
                elif line.strip() == '---' or line.strip() == '***':
                    story.append(Spacer(1, 10))
                # 图片处理
                elif line.strip().startswith('!['):
                    match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line.strip())
                    if match:
                        alt_text = match.group(1)
                        image_path = match.group(2)
                        if os.path.exists(image_path):
                            try:
                                from reportlab.platypus import Image as RLImage
                                img = RLImage(image_path)
                                # 限制图片宽度
                                img_width = 12 * cm  # 最大宽度
                                aspect = img.imageHeight / img.imageWidth
                                img_height = img_width * aspect
                                img.drawWidth = img_width
                                img.drawHeight = img_height
                                story.append(img)
                                if alt_text:
                                    story.append(Paragraph(f"<i>{alt_text}</i>", ParagraphStyle('caption', fontName=font_name, fontSize=8, alignment=TA_CENTER, textColor=colors.grey)))
                                story.append(Spacer(1, 6))
                            except Exception as img_err:
                                logger.warning(f"PDF 添加图片失败: {img_err}")
                                story.append(Paragraph(f"[图片: {alt_text}]", styles['body']))
                        else:
                            story.append(Paragraph(f"[图片不存在: {image_path}]", styles['body']))
                # 普通段落
                elif line.strip():
                    # 处理基础格式
                    text = line.replace('**', '').replace('*', '').replace('`', '')
                    if text.strip():
                        story.append(Paragraph(text, styles['body']))
                else:
                    story.append(Spacer(1, 4))

            # 生成 PDF
            filename = f"{doc_id}_{uuid.uuid4().hex[:8]}.pdf"
            filepath = os.path.join(self.export_dir, filename)

            doc = SimpleDocTemplate(filepath, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
            doc.build(story)

            return {
                "success": True,
                "format": ExportFormat.PDF,
                "filename": filename,
                "filepath": filepath,
                "size": os.path.getsize(filepath)
            }

        except ImportError as e:
            logger.error(f"缺少依赖库: {e}")
            return {
                "success": False,
                "error": f"缺少依赖库: {str(e)}",
                "format": ExportFormat.PDF
            }
        except Exception as e:
            logger.error(f"导出 PDF 失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "format": ExportFormat.PDF
            }

    def export_to_html(self, doc_id: str, md_content: str) -> Dict:
        """
        将 MD 转换为 HTML

        Args:
            doc_id: 文档ID
            md_content: Markdown 内容

        Returns:
            包含导出结果的字典
        """
        try:
            import markdown

            # 解析 Markdown 为 HTML
            html_body = markdown.markdown(
                md_content,
                extensions=['tables', 'fenced_code', 'codehilite', 'nl2br', 'toc']
            )

            # 构建完整 HTML
            html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <style>
        body {{
            font-family: "Microsoft YaHei", "SimSun", Arial, sans-serif;
            font-size: 14px;
            line-height: 1.8;
            max-width: 900px;
            margin: 0 auto;
            padding: 20px;
            color: #333;
        }}
        h1 {{ font-size: 28px; color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ font-size: 22px; color: #34495e; margin-top: 25px; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
        h3 {{ font-size: 18px; color: #555; margin-top: 20px; }}
        p {{ margin: 12px 0; }}
        table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
        th {{ background-color: #f8f9fa; font-weight: bold; }}
        code {{ background-color: #f5f5f5; padding: 2px 6px; border-radius: 3px; font-family: Consolas, "Courier New", monospace; font-size: 13px; }}
        pre {{ background-color: #f8f8f8; padding: 15px; border-radius: 5px; overflow-x: auto; border: 1px solid #e1e1e1; }}
        pre code {{ background: none; padding: 0; }}
        blockquote {{ border-left: 4px solid #3498db; margin: 15px 0; padding: 10px 15px; background-color: #f8f9fa; color: #555; }}
        ul, ol {{ margin: 10px 0; padding-left: 30px; }}
        li {{ margin: 6px 0; }}
        img {{ max-width: 100%; height: auto; border-radius: 5px; box-shadow: 0 2px 5px rgba(0,0,0,0.1); }}
        a {{ color: #3498db; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
        hr {{ border: none; border-top: 1px solid #eee; margin: 20px 0; }}
    </style>
</head>
<body>
{html_body}
</body>
</html>"""

            # 生成文件路径
            filename = f"{doc_id}_{uuid.uuid4().hex[:8]}.html"
            filepath = os.path.join(self.export_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(html_content)

            return {
                "success": True,
                "format": ExportFormat.HTML,
                "filename": filename,
                "filepath": filepath,
                "size": os.path.getsize(filepath)
            }

        except ImportError as e:
            logger.error(f"缺少依赖库: {e}")
            return {
                "success": False,
                "error": f"缺少依赖库: {str(e)}",
                "format": ExportFormat.HTML
            }
        except Exception as e:
            logger.error(f"导出 HTML 失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "format": ExportFormat.HTML
            }

    def export_to_txt(self, doc_id: str, md_content: str) -> Dict:
        """
        将 MD 转换为纯文本

        Args:
            doc_id: 文档ID
            md_content: Markdown 内容

        Returns:
            包含导出结果的字典
        """
        try:
            import re

            # 移除 Markdown 语法，保留纯文本
            text = md_content

            # 移除图片
            text = re.sub(r'!\[.*?\]\(.*?\)', '', text)

            # 移除链接，保留文字
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)

            # 移除代码块标记
            text = re.sub(r'```.*?\n', '', text)
            text = re.sub(r'`(.*?)`', r'\1', text)

            # 处理标题层级
            text = re.sub(r'^#### (.*)$', r'    \1', text, flags=re.MULTILINE)
            text = re.sub(r'^### (.*)$', r'  \1', text, flags=re.MULTILINE)
            text = re.sub(r'^## (.*)$', r'\1', text, flags=re.MULTILINE)
            text = re.sub(r'^# (.*)$', r'\1', text, flags=re.MULTILINE)

            # 处理加粗和斜体
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'__(.*?)__', r'\1', text)
            text = re.sub(r'_(.*?)_', r'\1', text)

            # 处理列表
            text = re.sub(r'^[\*\-] (.*)$', r'- \1', text, flags=re.MULTILINE)
            text = re.sub(r'^\d+\. (.*)$', r'\1', text, flags=re.MULTILINE)

            # 处理引用
            text = re.sub(r'^> (.*)$', r'  \1', text, flags=re.MULTILINE)

            # 移除多余空行
            text = re.sub(r'\n{3,}', '\n\n', text)

            # 生成文件路径
            filename = f"{doc_id}_{uuid.uuid4().hex[:8]}.txt"
            filepath = os.path.join(self.export_dir, filename)

            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(text)

            return {
                "success": True,
                "format": ExportFormat.TXT,
                "filename": filename,
                "filepath": filepath,
                "size": os.path.getsize(filepath)
            }

        except Exception as e:
            logger.error(f"导出 TXT 失败: {e}")
            return {
                "success": False,
                "error": str(e),
                "format": ExportFormat.TXT
            }

    def _add_html_content_to_doc(self, doc, html_content: str):
        """将 HTML 内容添加到 Word 文档"""
        from docx import Document
        from docx.shared import Pt
        from lxml import etree

        # 解析 HTML
        root = etree.fromstring(f"<div>{html_content}</div>", etree.HTMLParser())

        def process_element(element, doc):
            """递归处理 HTML 元素"""
            for child in element:
                tag = child.tag.lower() if isinstance(child.tag, str) else ''

                if tag == 'h1':
                    p = doc.add_heading(child.text or '', level=1)
                elif tag == 'h2':
                    p = doc.add_heading(child.text or '', level=2)
                elif tag == 'h3':
                    p = doc.add_heading(child.text or '', level=3)
                elif tag == 'h4':
                    p = doc.add_heading(child.text or '', level=4)
                elif tag == 'p':
                    text = ''.join(child.itertext())
                    if text.strip():
                        doc.add_paragraph(text)
                elif tag == 'ul':
                    for li in child:
                        if li.tag.lower() == 'li':
                            text = ''.join(li.itertext())
                            doc.add_paragraph(text, style='List Bullet')
                elif tag == 'ol':
                    for li in child:
                        if li.tag.lower() == 'li':
                            text = ''.join(li.itertext())
                            doc.add_paragraph(text, style='List Number')
                elif tag == 'blockquote':
                    text = ''.join(child.itertext())
                    p = doc.add_paragraph(text)
                    p_format = p.paragraph_format
                    p_format.left_indent = Inches(0.3)
                elif tag == 'pre':
                    text = ''.join(child.itertext())
                    p = doc.add_paragraph(text)
                    run = p.runs[0] if p.runs else p.add_run(text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                elif tag == 'code':
                    text = ''.join(child.itertext())
                    p = doc.add_paragraph(text)
                elif tag == 'table':
                    # 简单处理表格
                    rows = child.findall('.//tr')
                    if rows:
                        # 获取表头
                        headers = []
                        ths = rows[0].findall('.//th')
                        if ths:
                            headers = [th.text or '' for th in ths]
                        else:
                            tds = rows[0].findall('.//td')
                            headers = [td.text or '' for td in tds]

                        # 添加表头
                        if headers:
                            table = doc.add_table(rows=1, cols=len(headers))
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            for i, h in enumerate(headers):
                                hdr_cells[i].text = h

                        # 添加数据行
                        for row in rows[1:]:
                            tds = row.findall('.//td')
                            if tds:
                                row_cells = table.add_row().cells
                                for i, td in enumerate(tds[:len(headers)]):
                                    row_cells[i].text = td.text or ''
                elif tag == 'div' or tag == 'body':
                    process_element(child, doc)

        process_element(root, doc)


# 全局导出器实例
_exporter: Optional[MarkdownExporter] = None


def get_exporter() -> MarkdownExporter:
    """获取导出器单例"""
    global _exporter
    if _exporter is None:
        _exporter = MarkdownExporter()
    return _exporter


def export_md(doc_id: str, md_content: str, format: str) -> Dict:
    """
    导出 MD 为指定格式

    Args:
        doc_id: 文档ID
        md_content: Markdown 内容
        format: 目标格式 (docx, pdf, html, txt)

    Returns:
        导出结果
    """
    exporter = get_exporter()

    if format == ExportFormat.DOCX:
        return exporter.export_to_docx(doc_id, md_content)
    elif format == ExportFormat.PDF:
        return exporter.export_to_pdf(doc_id, md_content)
    elif format == ExportFormat.HTML:
        return exporter.export_to_html(doc_id, md_content)
    elif format == ExportFormat.TXT:
        return exporter.export_to_txt(doc_id, md_content)
    else:
        return {
            "success": False,
            "error": f"不支持的格式: {format}",
            "format": format
        }
